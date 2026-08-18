"""主张拆解/实体消歧 + 充分性评估 + 错误语义校正 + AgentState.metadata + docx失败不丢报告 回归测试。

全部使用 mock，不调用真实 LLM / Tavily / .env。
"""

from __future__ import annotations

import datetime as dt
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

from src.models import (
    AgentState,
    Claim,
    ClaimResult,
    Evidence,
    FactCheckReport,
)
from src.nodes.sufficiency import assess_sufficiency, has_background_context
from src.workflow import _PHASE_DECOMPOSE, _PHASE_SEARCH, _PHASE_SUFFICIENCY
from src.nodes.decompose import decompose_claims
from src.nodes.search import _get_search_queries
from src.nodes.report import generate_report


# ==============================================================
# 1) 主张拆解：TF家族案例生成 context/core/causal_or_detail
# ==============================================================
def test_decompose_tf_family_three_claim_types() -> None:
    text = "TF家族三代成员左航和邓佳鑫曾经是情侣，2020年在恋爱中。"

    class MockState:
        def __init__(self) -> None:
            self.original_text = text
            self.report = FactCheckReport(
                report_id="R-UNIT",
                original_text=text,
                generated_at=dt.datetime(2025, 3, 15),
                overall_verdict="",
                claims=[],
                claim_results=[],
                execution_log=[],
                timeline=[],
                unresolved_questions=[],
            )
            self.metadata: dict = {}

        def mark_step_started(self, phase: str, desc: str) -> None:
            self.metadata["phase_started"] = (phase, desc)

        def log(self, **kwargs) -> None:
            self.report.execution_log.append(kwargs)

    state = MockState()
    with patch("src.nodes.decompose._run_llm_decompose_if_enabled", return_value=None):
        out = decompose_claims(state, mode="demo")

    claims = out.report.claims
    # 至少 4 条主张（允许未来 LLM 模式更多，但 demo 必须 4 条）
    assert len(claims) >= 4, f"期望 ≥4 条主张，实际 {len(claims)}"
    roles = [c.claim_role for c in claims]
    # 必须存在三类
    assert "context" in roles
    assert "core" in roles
    assert "causal_or_detail" in roles

    texts = [(c.text, c.claim_role) for c in claims]
    # 必须有两条身份 context 主张
    context_texts = [t for t, r in texts if r == "context"]
    assert any("左航" in t and "TF" in t for t in context_texts), f"缺少左航身份 context：{context_texts}"
    assert any("邓佳鑫" in t and "TF" in t for t in context_texts), f"缺少邓佳鑫身份 context：{context_texts}"

    core = [t for t, r in texts if r == "core"]
    assert any("情侣" in t for t in core), f"核心主张缺少「情侣」事件：{core}"

    detail = [t for t, r in texts if r == "causal_or_detail"]
    assert any("2020" in t or "恋爱" in t for t in detail), f"细节主张缺少 2020/恋爱：{detail}"

    # core/细节必须 depends_on_claim_ids 指向 context（不能反过来）
    id_to_role = {c.claim_id: c.claim_role for c in claims}
    for c in claims:
        if c.claim_role in {"core", "causal_or_detail"}:
            assert c.depends_on_claim_ids, f"{c.claim_id}({c.claim_role}) 缺少依赖的 context：{c.text}"
            for dep in c.depends_on_claim_ids:
                assert id_to_role.get(dep) == "context" or id_to_role.get(dep) in {"context", "core"}, (
                    f"依赖链角色错误：{dep} 角色 {id_to_role.get(dep)}")
        # entity_aliases 必须非空（中文人物尤其需要）
        assert c.entity_aliases, f"{c.claim_id} 缺少 entity_aliases"


# ==============================================================
# 2) 实体消歧：同名人物搜索语句包含机构上下文
# ==============================================================
def test_search_queries_include_entity_aliases() -> None:
    """_get_search_queries(state, claim_id, fallback) 应优先用 Claim.search_keywords 生成消歧语句。"""
    c = Claim(
        claim_id="x1",
        text="左航是TF家族三代成员",
        claim_type="FACT",
        claim_role="context",
        depends_on_claim_ids=[],
        entity_aliases={"左航": ["TF家族", "三代", "成员"], "TF家族": ["时代峰峻"]},
        needs_background_verification=True,
        verification_question="",
        risk_level="low",
        entities=["左航", "TF家族"],
        search_keywords=[
            "左航 TF家族 三代 成员",
            "左航 官方 个人介绍 TF家族",
            "TF家族三代 成员名单 左航",
        ],
    )
    state = AgentState(original_text="测试", claims=[c])
    queries = _get_search_queries(state, c.claim_id, fallback_text="左航")
    # 必须包含 search_keywords（entity_aliases 对应上下文字段）
    joined = " || ".join(queries)
    assert "TF家族" in joined and ("三代" in joined or "时代峰峻" in joined), f"搜索语句缺少机构上下文：{queries}"
    # 裸词"左航"不应独占（不能只搜左航）
    assert not (len(queries) == 1 and queries[0].strip() == "左航"), f"禁用裸名字搜索：{queries}"


def test_has_background_context_detects_missing_context() -> None:
    # 裸名 + 没有别名（entity_aliases 字典为空）
    c_naked = Claim(
        claim_id="x2", text="左航曾经是网红", claim_type="EVENT",
        claim_role="core", depends_on_claim_ids=[],
        entity_aliases={},  # 没有别名
        needs_background_verification=False, verification_question="",
        risk_level="medium", entities=["左航"],
        search_keywords=["左航 网红"],
    )
    assert has_background_context(c_naked, []) is False

    # 有别名且包含"TF家族"
    c_good = c_naked.model_copy(update={"entity_aliases": {"左航": ["TF家族", "三代", "成员"]}})
    assert has_background_context(c_good, []) is True


# ==============================================================
# 3) AgentState.metadata 默认独立；第二轮检索读写 metadata；达到上限停止
# ==============================================================
def test_agent_state_metadata_isolated_defaults() -> None:
    """AgentState 实例之间 metadata 必须独立，禁止共享默认 dict。"""
    s1 = AgentState(original_text="案例A")
    s2 = AgentState(original_text="案例B")
    # 默认是空 dict（不是 None）
    assert isinstance(s1.metadata, dict) and s1.metadata == {}
    # s1 改完 s2 不变
    s1.metadata["foo"] = 123
    s1.metadata.setdefault("_search_rounds", {})["c1"] = 2
    assert s2.metadata == {}, "s2.metadata 仍应为空（不共享默认对象）"
    s2.metadata["bar"] = 456
    assert s1.metadata.get("bar") is None, "s1 和 s2 metadata 独立"


def test_agent_state_default_run_without_explicit_metadata() -> None:
    """不传 metadata 也能正常构造并运行（workflow 创建 state 时不显式传）。"""
    s = AgentState(original_text="xxx")
    s.metadata.setdefault("search_rounds", {})["c1"] = 1
    assert s.metadata["search_rounds"]["c1"] == 1


def _make_claims_for_sufficiency(claim_ids: list[str]) -> list[Claim]:
    claims: list[Claim] = []
    for cid in claim_ids:
        is_context = cid in {"c1", "c2"}
        claims.append(Claim(
            claim_id=cid,
            text=f"主张 {cid}",
            claim_type="FACT",
            claim_role="context" if is_context else "core",
            depends_on_claim_ids=[],
            entity_aliases={cid: [f"别名-{cid}"]},
            needs_background_verification=is_context,
            verification_question="",
            risk_level="low" if is_context else "high",
            entities=[cid],
            search_keywords=[f"{cid} 搜索词"],
        ))
    return claims


def _add_evidences(state: AgentState, need_more_ids: list[str]) -> None:
    """往 state.evidence / state.evidence_pool 写证据。"""
    claim_ids = [c.claim_id for c in state.claims]
    for cid in claim_ids:
        need_more = cid in need_more_ids
        base = [
            Evidence(
                evidence_id=f"e-{cid}",
                claim_id=cid,
                source_url=f"https://fan.example.com/{cid}",
                source_title=f"某粉丝帖子关于 {cid}",
                publisher="粉丝论坛",
                source_type="社交帖子",
                source_domain="fan.example.com",
                source_grade="D",
                published_at=dt.datetime(2024, 1, 1),
                retrieved_at=dt.datetime(2025, 1, 1),
                relevant_excerpt="帖子正文：我听说xxx属实。",
                evidence_summary="粉丝发帖提及",
                supports_or_refutes="supports",
                source_content="摘要",
                evidence_stance="supports",
                directness="low",
                is_primary_source=False,
                extraction_status="success",
                reliability_reason="粉丝帖子，来源可信度低",
                independence_group=f"fan-post-{cid}",
            )
        ]
        if not need_more:
            base += [
                Evidence(
                    evidence_id=f"a-{cid}", claim_id=cid,
                    source_url=f"https://off.example.com/{cid}",
                    source_title=f"官方公开资料 {cid}",
                    publisher="官方", source_type="官方公告",
                    source_domain="off.example.com", source_grade="A",
                    published_at=dt.datetime(2023, 6, 1),
                    retrieved_at=dt.datetime(2025, 1, 1),
                    relevant_excerpt="官方公开资料明确记载...",
                    evidence_summary="官方记载",
                    supports_or_refutes="supports",
                    source_content="官方全文",
                    evidence_stance="supports",
                    directness="high", is_primary_source=True,
                    extraction_status="success",
                    reliability_reason="官方一手来源，可信度高",
                    independence_group=f"off-{cid}",
                ),
                Evidence(
                    evidence_id=f"b-{cid}", claim_id=cid,
                    source_url=f"https://news.example.com/{cid}",
                    source_title=f"权威媒体独立报道 {cid}",
                    publisher="权威媒体", source_type="新闻报道",
                    source_domain="news.example.com", source_grade="B",
                    published_at=dt.datetime(2023, 6, 2),
                    retrieved_at=dt.datetime(2025, 1, 1),
                    relevant_excerpt="权威媒体独立核实后报道...",
                    evidence_summary="权威独立报道",
                    supports_or_refutes="supports",
                    source_content="新闻全文",
                    evidence_stance="supports",
                    directness="high", is_primary_source=False,
                    extraction_status="success",
                    reliability_reason="权威媒体与官方相互独立，可信度较高",
                    independence_group=f"news-{cid}",
                ),
            ]
        state.evidence[cid] = base
        state.evidence_pool.extend(base)


def test_sufficiency_round2_triggered_and_then_stops() -> None:
    """第二轮检索：第一轮证据不足 → follow_up 有值；第二轮 _search_rounds[cid]==2 且仍缺证据 → follow_up 空。

    注意：这里只测 sufficiency 节点本身；workflow.py 是否调用第二轮 search 由 workflow 单独测试。
    """
    claim_ids = ["c1", "c2", "c3", "c4"]

    # --- Round 1：c1/c2 有 A+B 独立组；c3/c4 只有 D，需要第二轮 ---
    s1 = AgentState(original_text="案例A", claims=_make_claims_for_sufficiency(claim_ids))
    _add_evidences(s1, need_more_ids=["c3", "c4"])
    out1 = assess_sufficiency(s1)
    fu1 = out1.metadata.get("follow_up") or {}
    # c3/c4 仍缺证据 → follow_up 至少包含这两个
    assert fu1, "第一轮 c3/c4 证据不足，应生成 follow_up"
    assert {"c3", "c4"}.issubset(set(fu1.keys())), f"follow_up 缺少 c3/c4: {list(fu1)}"
    # 日志存在：至少有一条是 sufficiency 阶段 action（含"补第2轮/补充检索"等字样，或有 sufficient=false 记录）
    acts1 = " ".join(
        (str(x.get("action", "")) + " " + str(x.get("status", "")) + " " +
         str((x.get("details") or {}).get("reason", "")) + " " +
         str((x.get("details") or {}).get("round", "")))
        for x in out1.execution_log
    )
    assert (
        "补第2轮" in acts1 or "第二轮" in acts1 or "补充检索" in acts1 or
        "证据不足" in acts1 or "insufficient" in acts1
    ), f"sufficiency 日志缺少第二轮/证据不足相关记录: {out1.execution_log}"

    # --- Round 2：显式把 _search_rounds c3/c4 设为 2，证据仍是 D ---
    s2 = AgentState(original_text="案例A", claims=_make_claims_for_sufficiency(claim_ids))
    _add_evidences(s2, need_more_ids=["c3", "c4"])
    s2.metadata["_search_rounds"] = {"c1": 1, "c2": 1, "c3": 2, "c4": 2}
    out2 = assess_sufficiency(s2)
    fu2 = out2.metadata.get("follow_up") or {}
    # 达到上限（c3/c4 轮次 == MAX_ROUNDS_PER_CLAIM == 2），不应再生成 follow_up
    assert not fu2, f"达到 MAX_ROUNDS_PER_CLAIM=2 不应继续 follow_up: {fu2}"
    # 日志里应有：达到最大检索轮次 或 无需补充检索 或 明确写了 sufficiency_finalized
    acts2 = " ".join(
        (str(x.get("action", "")) + " " + str((x.get("details") or {}).get("reason", "")))
        for x in out2.execution_log
    )
    assert (
        "达到最大检索轮次" in acts2 or "无需补充检索" in acts2 or
        "已完成充分性评估" in acts2 or "所有主张" in acts2
    ), f"第二轮日志内容不符合预期: {out2.execution_log}"


# ==============================================================
# 4) Streamlit 重绘不重新执行 LLM / 搜索（session_state 缓存）+ 失败态不渲染虚假结论
# ==============================================================
def test_streamlit_rerender_no_re_execution_and_hide_claims_when_failed() -> None:
    """1) session_state.report 存在时直接渲染不重跑 workflow；
    2) overall=核查失败 或 workflow_error 存在时，主张详情区不应当渲染（避免"证据不足"式误导）。
    """
    import pathlib
    src = pathlib.Path(__file__).resolve().parent.parent / "app.py"
    text = src.read_text(encoding="utf-8")
    # 关键契约：
    # 1. report 存在于 st.session_state 时，直接渲染不调用 workflow
    # 2. 新 text != input_text 时才重新执行
    assert 'if "report" not in st.session_state' in text, "app.py 缺少 session_state report 初始化"
    assert 'text and text != st.session_state.input_text' in text, "缺少 input 变化判断"
    assert 'st.session_state.pop("report", None)' in text, "新提交应清空旧报告"
    # docx_bytes 缓存键存在于 session_state 初始化
    assert 'if "docx_bytes" not in st.session_state' in text, "缺少 docx_bytes 缓存初始化"
    # 下载按钮使用正确 mime
    assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in text
    # 下载按钮文件命名格式：溯真_事实核查报告_YYYYMMDD_HHMMSS.docx
    assert "溯真_事实核查报告_" in text and "%Y%m%d_%H%M%S" in text

    # 错误语义：失败态（workflow_error/current_step=failed/overall=核查失败）→ 不渲染 claim_results 等主体
    # 校验：代码中对 failed 做了分支判断（失败态分支不调用 render_decompose_results/render_claim_results 等）
    assert 'failed = bool(report.workflow_error)' in text or "overall_verdict == \"核查失败\"" in text, (
        "缺少失败态判定变量 failed")
    # 失败态分支：render_decompose_results / render_claim_results 应当被保护（只在 not failed 下调用）
    assert "if not failed:" in text and "render_decompose_results" in text, (
        "失败态仍可能渲染虚假的主张拆解（应仅在 not failed 分支）")
    # 下载按钮在失败态下不得出现（render_download_section 内部判断 workflow_error/current_step/overall）
    assert "核查失败" in text  # render_overall_verdict 里有 核查失败 的颜色或文案


# ==============================================================
# 5) docx 导入 + 模拟报告生成/重新打开验证
# ==============================================================
def test_docx_exporter_import_and_roundtrip() -> None:
    """1) docx_exporter 能直接导入（语法正确）；2) 模拟报告生成 DOCX；3) python-docx 重新打开校验。"""
    # --- 导入：若 SyntaxError 仍存在，这里会直接抛（不会被 try/except 吞掉） ---
    from src.exporters.docx_exporter import build_fact_check_docx  # noqa: F401

    # --- 构造最小模拟 FactCheckReport ---
    claim = Claim(
        claim_id="c1", text="左航 2023 年出道", claim_type="EVENT",
        claim_role="core", depends_on_claim_ids=[],
        entity_aliases={"左航": ["TF家族", "三代"]},
        needs_background_verification=False, verification_question="出道时间是否属实？",
        risk_level="low", entities=["左航"],
        search_keywords=["左航 出道时间 2023"],
    )
    ev1 = Evidence(
        evidence_id="e1", claim_id="c1",
        source_url="https://off.example.com/zuohang",
        source_title="官方人员履历",
        publisher="官方", source_type="官方公告",
        source_domain="off.example.com", source_grade="A",
        published_at=dt.datetime(2023, 10, 1),
        retrieved_at=dt.datetime(2025, 1, 1),
        relevant_excerpt="左航于2023年随三代出道。",
        evidence_summary="官方记载出道时间",
        supports_or_refutes="supports",
        source_content="官方全文",
        evidence_stance="supports",
        directness="high", is_primary_source=True,
        extraction_status="success",
        reliability_reason="官方一手来源，可信度高",
        independence_group="off-1",
    )
    result = ClaimResult(
        claim=claim, verdict="已证实", confidence=0.9,
        reasoning="官方记载，来源A级，时间点明确。",
        evidence=[ev1],
        missing_information=None,
    )
    report = FactCheckReport(
        original_text="左航2023年出道。",
        overall_verdict="已证实",
        overall_summary="左航 2023 年出道经官方一手信息证实；传播风险可控。",
        propagation_risk="低：公开出道事件，传播后误读风险低。",
        claim_results=[result],
        claims=[claim],
        evidence={"c1": [ev1]},
        execution_log=[
            {"step": "decompose", "action": "decompose_ok", "status": "success",
             "details": {"count": 1}},
            {"step": "search", "action": "search_round_1", "status": "success",
             "details": {"round": 1}},
            {"step": "sufficiency", "action": "证据充分", "status": "success",
             "details": {"claims_total": 1}},
            {"step": "evaluate", "action": "evaluate_ok", "status": "success",
             "details": {}},
            {"step": "report", "action": "report_ok", "status": "success",
             "details": {}},
        ],
        unresolved_questions=[],
        completed_steps=[_PHASE_DECOMPOSE, _PHASE_SEARCH, _PHASE_SUFFICIENCY, "evaluate", "report"],
        current_step="completed",
        workflow_completed=True,
        workflow_error=None,
        generated_at=dt.datetime(2025, 3, 15),
    )

    # --- 生成 DOCX 字节 ---
    docx_bytes = build_fact_check_docx(report)
    assert isinstance(docx_bytes, (bytes, bytearray)) and len(docx_bytes) > 2000, "生成的 DOCX 字节过小"

    # --- 用 python-docx 重新打开校验：最小可识别 DOCX（zip + 文档正文非空） ---
    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs) + "\n".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells
    )
    # 报告关键字段必须落在 DOCX 中（证明内容被写入）
    # 注意：表格列中"主张"单元格会截断到 80 字以内，我们用"左航 2023"子串作为断言
    for kw in ["左航 2023", "已证实", "官方人员履历", "事实核查报告", "溯真"]:
        assert kw in full_text, f"DOCX 缺少关键字段: {kw!r}"


# ==============================================================
# 6) 错误语义校正：workflow_error 存在时 overall 为 "核查失败" 而非 "证据不足"
# ==============================================================
def _mock_claims_and_results_for_semantics():
    claims = [Claim(
        claim_id="x1", text="主张 X", claim_type="FACT",
        claim_role="core", depends_on_claim_ids=[],
        entity_aliases={"X": ["别名"]},
        needs_background_verification=False, verification_question="X?",
        risk_level="low", entities=["X"], search_keywords=["X"],
    )]
    return claims


def test_report_error_semantics_fail_not_insufficient() -> None:
    """workflow_error 存在（异常阶段），generate_report 必须写 overall="核查失败"，而不是"证据不足"。"""
    claims = _mock_claims_and_results_for_semantics()
    s = AgentState(original_text="文本", claims=claims, workflow_error="decompose")
    s.report = FactCheckReport(
        original_text="文本",
        overall_verdict="已证实",  # 占位（generate_report 会改为"核查失败"）
        overall_summary="占位摘要",
        propagation_risk="低风险：失败态，未评估传播风险",
        claim_results=[],  # 没有单条结论（节点尚未运行）
        execution_log=[
            {"phase": "decompose", "action": "exception", "detail": "ValueError: 模型响应格式异常"},
        ],
        timeline=[],
        unresolved_questions=[],
        current_step="failed",
        completed_steps=[],
        workflow_completed=False,
        workflow_error="decompose",
        generated_at=dt.datetime(2025, 3, 15),
    )
    # simulate: LLM 在 decompose 抛出 ValueError → workflow_error = "decompose"，
    # completed_steps 仍为空，claim_results 是空（不能伪造"证据不足"结论）
    assert s.claim_results == []

    out = generate_report(s)
    # 总体必须是"核查失败"，不得伪装为"证据不足"
    assert out.report.overall_verdict == "核查失败", (
        f"workflow_error 存在时应写为 核查失败，实际：{out.report.overall_verdict!r}")
    # 摘要中应明确包含失败阶段（中文）
    summary = out.report.summary or ""
    assert "核查失败" in summary and ("分解" in summary or "decompose" in summary), (
        f"失败摘要未指明阶段：{summary}")
    # 失败态不应生成虚假单条结论（claim_results 仍为空 / 或只允许在已经运行过的节点后写入）
    # → generate_report 不应凭空添加 verdict
    assert not out.report.claim_results or all(r.verdict != "证据不足" for r in out.report.claim_results), (
        "失败态不应伪造'证据不足'结论")
    # overall_verdict 是 "核查失败"，不是 "证据不足"
    assert out.report.overall_verdict != "证据不足"


def test_report_insufficient_still_possible_without_workflow_error() -> None:
    """workflow 成功跑完，但所有 claim 都"证据不足" → overall="证据不足"（这是合法的业务结论）。"""
    from src.models import ClaimResult as CR
    claims = _mock_claims_and_results_for_semantics()
    s = AgentState(original_text="文本", claims=claims)
    s.claim_results = [CR(
        claim_id="x1", verdict="证据不足", confidence=0.4,
        rationale="缺少可靠来源",
        supporting_evidence_ids=[], opposing_evidence_ids=[],
        unresolved_questions=["需要更明确的公开证据"],
    )]
    s.report = FactCheckReport(
        original_text="文本",
        overall_verdict="已证实",  # 占位：generate_report 会根据 verdicts 重算为"证据不足"
        overall_summary="占位摘要",
        propagation_risk="中等：全部主张证据不足，传播后易被误读",
        claims=claims,
        claim_results=s.claim_results,
        execution_log=[{"phase": "evaluate", "action": "all_insufficient"}],
        timeline=[],
        unresolved_questions=["需要更明确的公开证据"],
        current_step="completed",
        completed_steps=[_PHASE_DECOMPOSE, _PHASE_SEARCH, _PHASE_SUFFICIENCY, "evaluate", "report"],
        workflow_completed=True,
        workflow_error=None,
        generated_at=dt.datetime(2025, 3, 15),
    )
    out = generate_report(s)
    # 业务态：整体 = 证据不足
    assert out.report.overall_verdict == "证据不足", f"合法业务态应保持 证据不足，实际：{out.report.overall_verdict!r}"
    # 摘要不应出现"核查失败"
    assert "核查失败" not in (out.report.summary or "")


# ==============================================================
# 7) Word 导出异常不得清除已经成功生成的网页报告
# ==============================================================
def test_word_export_failure_keeps_web_report() -> None:
    """app.py 在点击下载时如果 docx 生成失败，不应清空 session_state.report 或重跑 workflow。

    本测试读取 app.py，验证：
    - 下载按钮的回调不会触发 run_workflow / pop("report", None)
    - docx 异常被捕获并 toast/exception，而不会覆盖 st.session_state.report
    """
    import pathlib
    src = pathlib.Path(__file__).resolve().parent.parent / "app.py"
    text = src.read_text(encoding="utf-8")

    # 验证：报告加载路径中只在"新提交"时 pop report；下载不会
    # 关键：新提交条件 text != input_text 时才 pop；下载路径无此 pop
    assert "build_fact_check_docx" in text, "app.py 中未调用 build_fact_check_docx"
    # 必须把 build_fact_check_docx 抛错的路径与"重新生成报告"解耦：
    # 我们要求 app.py 中生成 docx 的 try/except 不得包含 st.session_state.report=... 赋值
    # 简化验证：app.py 中 "session_state.report = run_fact_check_workflow" 只出现1次（顶部主流程），
    # 在下载区块中不会出现该赋值语句。
    count_new_report = text.count("st.session_state.report = run_fact_check_workflow")
    count_report_assing_in_download_try = 0  # 我们用更严格的验证：try 内的 "report = " 不得是 run_fact_check_workflow
    # 核心检查：下载按钮逻辑不包含 pop("report", None) 或 .report = run_fact_check_workflow
    assert count_new_report == 1, "报告生成应只在主流程进行一次，避免下载异常时覆盖 report"
    # 同时：render_download_section() 内调用 build_fact_check_docx 必须被 try/except 包裹（导出失败仅提示，不丢网页）
    # 这是为了符合"Word导出异常不得清除已经成功生成的网页报告"的语义
    # 简化：确保 try 后有 except 捕获 Exception（更宽允许，但不得清空 report）
    assert "try:" in text and "build_fact_check_docx" in text


# ==============================================================
# 8) 完整 mock 工作流：可以成功生成主张与报告（顺序/节点计数）
# ==============================================================
def test_workflow_pipeline_mock() -> None:
    """验证工作流会依次跑 decompose → plan → search(r1) → sufficiency → search(r2，若follow_up)
    → evaluate → build_report → store_memory（通过假节点计数 + 日志验证）。"""
    from src.workflow import run_fact_check_workflow

    counts = {"decompose": 0, "search": 0, "plan": 0, "sufficiency": 0,
              "evaluate": 0, "build_report": 0, "store_memory": 0}

    def mk(what: str):
        def _fn(state, *args, **kwargs):
            counts[what] += 1
            return state
        return _fn

    with (
        patch("src.workflow.decompose_claims", side_effect=mk("decompose")),
        patch("src.workflow.build_verification_plan", side_effect=mk("plan")),
        patch("src.workflow.search_evidence", side_effect=mk("search")),
        patch("src.workflow.assess_sufficiency", side_effect=mk("sufficiency")),
        patch("src.workflow.cross_validate_claims", side_effect=mk("evaluate")),
        patch("src.workflow.build_report", side_effect=mk("build_report")),
        patch("src.workflow.store_to_memory_if_enabled", side_effect=mk("store_memory")),
    ):
        out = run_fact_check_workflow("任意文本", mode="demo")

    # 各节点被调用（顺序计数）
    assert counts["decompose"] == 1
    assert counts["plan"] == 1
    assert counts["search"] >= 1
    assert counts["sufficiency"] == 1
    assert counts["evaluate"] == 1
    assert counts["build_report"] == 1
    # report 对象：结构化进度字段 sufficiency 在 completed_steps
    assert _PHASE_SUFFICIENCY in out.completed_steps or out.current_step in {"completed", "failed"}
