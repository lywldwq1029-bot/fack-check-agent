"""DOCX 导出回归测试（全部使用模拟数据，不调用真实 API、不读密钥）。"""

from __future__ import annotations

import datetime as dt
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from src.exporters.docx_exporter import build_fact_check_docx
from src.models import (
    Claim,
    ClaimResult,
    Evidence,
    FactCheckReport,
    TimelineEvent,
    VerificationPlan,
)


# ========================= fixtures =========================
def _make_report(include_timeline: bool = True) -> FactCheckReport:
    """构造一份覆盖所有分支的模拟报告（严格按 src/models.py 中 FactCheckReport.model_fields 构建）。"""
    c1 = Claim(
        claim_id="c1",
        text="左航是TF家族三代成员，男团养成计划公开练习生。" + ("A" * 80),  # 长标题
        claim_type="FACT",
        claim_role="context",
        depends_on_claim_ids=[],
        entity_aliases={"左航": ["TF家族", "三代", "成员"]},
        needs_background_verification=True,
        verification_question="左航是否为TF家族三代成员？",
        risk_level="low",
        entities=["左航", "TF家族"],
        search_keywords=["左航 TF家族 三代 成员", "左航 官方 个人介绍 TF家族"],
    )
    c2 = Claim(
        claim_id="c2",
        text="邓佳鑫是TF家族三代成员。",
        claim_type="FACT",
        claim_role="context",
        depends_on_claim_ids=[],
        entity_aliases={"邓佳鑫": ["TF家族", "三代", "成员"]},
        needs_background_verification=True,
        verification_question="邓佳鑫是否为TF家族三代成员？",
        risk_level="low",
        entities=["邓佳鑫", "TF家族"],
        search_keywords=["邓佳鑫 TF家族 三代 成员"],
    )
    c3 = Claim(
        claim_id="c3",
        text="左航和邓佳鑫曾经是情侣关系。",
        claim_type="EVENT",
        claim_role="core",
        depends_on_claim_ids=["c1", "c2"],
        entity_aliases={"左航": ["邓佳鑫", "TF家族", "三代"]},
        needs_background_verification=False,
        verification_question="左航与邓佳鑫是否曾经是情侣？",
        risk_level="high",
        entities=["左航", "邓佳鑫", "情侣"],
        search_keywords=["左航 邓佳鑫 情侣 TF家族", "左航 邓佳鑫 回应 关系"],
    )
    c4 = Claim(
        claim_id="c4",
        text="两人在2020年处于恋爱关系中。",
        claim_type="EVENT",
        claim_role="causal_or_detail",
        depends_on_claim_ids=["c3"],
        entity_aliases={"左航": ["邓佳鑫", "2020", "恋爱"]},
        needs_background_verification=False,
        verification_question="左航与邓佳鑫是否在2020年处于恋爱关系？",
        risk_level="high",
        entities=["左航", "邓佳鑫", "2020年"],
        search_keywords=["左航 邓佳鑫 2020 恋爱"],
    )

    claims = [c1, c2, c3, c4]

    def ev(eid: str, url: str, title: str, publisher: str, grade: str, stance: str,
           excerpt: str, published_at: Optional[dt.datetime], source_type: str = "新闻报道",
           domain: str = "", group: str = "") -> Evidence:
        return Evidence(
            evidence_id=eid,
            claim_id=(
                {"e1": "c1", "e2": "c2", "e3": "c1", "e4": "c3", "e5": "c3",
                 "e6": "c3", "e7": "c4", "e8": "c3"}.get(eid, "c1")
            ),
            source_url=url,
            source_title=title,
            publisher=publisher,
            source_type=source_type,
            source_domain=domain or publisher,
            source_grade=grade,
            published_at=published_at,
            retrieved_at=dt.datetime(2025, 1, 1, 12, 0, 0),
            relevant_excerpt=excerpt,
            evidence_summary=excerpt[:60],
            supports_or_refutes=("supports" if stance == "supports" else "refutes"),
            is_primary_source=grade in {"A", "B"},
            reliability_reason=(
                "官方一手来源，可信度高" if grade in {"A"}
                else "第三方权威独立媒体，可信度较高" if grade == "B"
                else "粉丝匿名发帖/转载，可信度低"
            ),
            source_content=excerpt,
            evidence_stance=stance,
            directness="high",
            extraction_status="success",
            independence_group=group or domain,
        )

    now = dt.datetime(2025, 3, 15, 10, 0, 0)
    e1 = ev("e1", "https://tf.example.com/zuohang",
            "TF家族官方练习生名录：三代训练生名单（更新至2024）",
            "TF家族官网", "A", "supports",
            "三代训练生名单如下：左航、陈天润、童禹坤……；本页为官方公开的三代训练生完整名录。",
            dt.datetime(2024, 1, 10, 12, 0), domain="tf.example.com", group="tf-official")
    e2 = ev("e2", "https://tf.example.com/dengjiaxin",
            "官方公告：关于邓佳鑫为TF家族三代公开训练生的说明",
            "TF家族官网", "A", "supports",
            "邓佳鑫为我司TF家族三代公开训练生，特此说明。",
            dt.datetime(2023, 6, 1, 10, 30), domain="tf.example.com", group="tf-official")
    e3 = ev("e3", "https://news.example.com/tf-3rd",
            "新京报：TF家族三代练习生名单首次公布",
            "新京报", "B", "supports",
            "时代峰峻旗下TF家族公开三代练习生名单，其中左航、邓佳鑫均在列。",
            dt.datetime(2021, 5, 20, 9, 0), domain="news.example.com", group="news-xjb")
    e4 = ev("e4", "https://weibo.example.com/fan-1",
            "某粉丝论坛：我嗑的CP一定是真的！（匿名）",
            "粉丝论坛", "D", "supports",
            "左航邓佳鑫2020年一定在一起了，你看那个眼神！（匿名用户发帖）",
            dt.datetime(2020, 12, 25, 22, 0), domain="weibo.example.com", group="fan-post-1")
    e5 = ev("e5", "https://blog.example.com/repost",
            "某自媒体：网传左航邓佳鑫曾是情侣（转载粉丝爆料）",
            "自媒体博客", "D", "supports",
            "据粉丝匿名爆料，左航邓佳鑫2020年曾在恋爱中。",
            dt.datetime(2021, 1, 5, 11, 0), domain="blog.example.com", group="fan-post-1")
    e6 = ev("e6", "https://self-media.example.com/repost2",
            "娱乐号：再次转载左航邓佳鑫恋爱传闻",
            "娱乐自媒体号", "D", "supports",
            "（与e5同源爆料）左航邓佳鑫2020年疑似恋爱。",
            dt.datetime(2021, 1, 6, 18, 0), domain="self-media.example.com", group="fan-post-1")
    e7 = ev("e7", "https://video.example.com/clip",
            "哔哩哔哩剪辑：某CP向剪辑视频标题【左邓2020】",
            "B站UP主剪辑", "E", "supports",
            "（粉丝二次创作剪辑视频标题）左航邓佳鑫2020同框剪辑合集。",
            dt.datetime(2022, 2, 14, 0, 0), domain="video.example.com", group="fan-clip")
    e8 = ev("e8", "https://news.example.com/deny",
            "当事人所属公司声明：关于不实传闻的澄清",
            "时代峰峻官方", "A", "refutes",
            "针对近期涉及我司艺人的不实私生活传闻，我司严正声明：相关说法均为杜撰。",
            dt.datetime(2021, 2, 1, 16, 0), domain="news.example.com", group="official-deny")

    evidences_flat = [e1, e2, e3, e4, e5, e6, e7, e8]
    # FactCheckReport.evidence 现在不是模型显式字段；通过 ClaimResult.evidence 传递给 DOCX
    # （docx_exporter 仍会读取 ctx.evidence_index，它会从 FactCheckReport.evidence 读取，但
    #  FactCheckReport 没有 evidence 字段——因此这里额外构造 evidence={cid:[...] 作为
    #  **extra（Pydantic 默认允许额外字段）给 docx_exporter 做兼容填充）

    # claim 结果：背景证实，私生活证据不足，时间细节证据不足
    r1 = ClaimResult(
        claim=c1, verdict="已证实", confidence=0.95,
        reasoning="TF家族官方公开名录、新京报独立报道均显示左航为三代训练生，"
                  "A/B 级多独立来源一致，判断为已证实。",
        evidence=[e1, e3],
        missing_information="",
    )
    r2 = ClaimResult(
        claim=c2, verdict="已证实", confidence=0.94,
        reasoning="TF家族官方声明 + 第三方独立媒体报道相互佐证，判断为已证实。",
        evidence=[e2, e3],
        missing_information="",
    )
    r3 = ClaimResult(
        claim=c3, verdict="证据不足", confidence=0.30,
        reasoning="该主张为自然人私生活高风险指控。当前所有支持证据均来自匿名粉丝发帖、"
                  "自媒体转载与粉丝剪辑视频，转载内容共用同一爆料按 1 个独立来源计算；"
                  "缺乏当事人公开声明、署名主流媒体采访或 ≥2 个相互独立的高质量来源。"
                  "同时存在所属公司的 A 级澄清声明。"
                  "因此结论为证据不足，不能判定属实。",
        evidence=[e4, e5, e6, e7, e8],
        missing_information="需要当事人公开回应、署名媒体的调查报道或双方任何一方的声明性陈述。",
    )
    r4 = ClaimResult(
        claim=c4, verdict="证据不足", confidence=0.25,
        reasoning="关于 2020 年恋爱关系，缺乏任何 A/B 级独立来源，"
                  "匿名爆料与转载不构成可靠证据链。结论为证据不足。",
        evidence=[e4, e5, e6],
        missing_information="需要 2020 年时间点的当事人声明或可靠媒体报道。",
    )

    claim_results = [r1, r2, r3, r4]

    # VerificationPlan：按真实 models.py 定义（单条主张维度；不是全局 case_summary 结构）
    # 这里给 c1/c2/c3/c4 各一个 plan 条目（list[VerificationPlan]）；FactCheckReport
    # 当前没有 plan 字段——以 **extra 方式填充，便于 docx_exporter 未来若增加
    # 渲染计划区块时不会因为字段名不存在而直接 AttributeError（目前 _render 计划
    # 节点不存在，docx_exporter 没有读取 plan，所以 extra 不影响正确性）
    plans_extra = {
        "plan": [
            VerificationPlan(
                claim_id=cid,
                verification_steps=steps,
                search_queries=queries,
                preferred_sources=["官方公开资料", "权威独立媒体"],
                required_evidence_level="一般证据",
                priority=1,
                priority_reason=reason,
            )
            for cid, steps, queries, reason in [
                ("c1", ["搜索官方公开名录与权威媒体报道"],
                 ["左航 TF家族 三代 成员", "左航 官方 个人介绍 TF家族"],
                 "人物身份类主张必须具备官方一手来源"),
                ("c2", ["搜索官方声明与权威媒体报道"],
                 ["邓佳鑫 TF家族 三代 成员"],
                 "人物身份类主张必须具备官方一手来源"),
                ("c3", ["搜索当事人声明、署名媒体报道、澄清公告；识别转载同源性"],
                 ["左航 邓佳鑫 情侣 TF家族", "左航 邓佳鑫 回应 关系"],
                 "私生活高风险主张需要 ≥2 个独立高质量来源"),
                ("c4", ["搜索 2020 年时间点的声明、报道"],
                 ["左航 邓佳鑫 2020 恋爱"],
                 "含具体时间点的高风险细节需要时间维度上的独立证据"),
            ]
        ]
    }

    timeline: list[TimelineEvent] = []
    if include_timeline:
        timeline = [
            TimelineEvent(
                event_time=dt.datetime(2020, 12, 25, 22, 0),
                description="匿名粉丝发帖提及传闻（2020-12-25，来源于 e4）",
                source_url="https://weibo.example.com/fan-1",
            ),
            TimelineEvent(
                event_time=dt.datetime(2021, 1, 5, 11, 0),
                description="自媒体转载匿名爆料（与 e4 同源，按 1 个独立来源计）",
                source_url="https://blog.example.com/repost",
            ),
            TimelineEvent(
                event_time=dt.datetime(2021, 2, 1, 16, 0),
                description="当事人所属公司发布 A 级澄清声明（e8）",
                source_url="https://news.example.com/deny",
            ),
            TimelineEvent(
                event_time=dt.datetime(2021, 5, 20, 9, 0),
                description="权威媒体公开三代训练生名单（e3）",
                source_url="https://news.example.com/tf-3rd",
            ),
        ]

    # propagation_risk：当前 FactCheckReport 字段为 str（非结构化对象）
    propagation_risk_str = (
        "高：该传闻涉及自然人私生活（恋爱关系）及具体时间维度，在粉丝圈层易被放大转载；"
        "一旦传播未证实内容，存在侵犯名誉权、误导公众的风险。"
        "建议用户关注当事人及所属公司的正式声明，避免转发未证实的匿名爆料、转载与剪辑内容。"
    )

    unresolved = [
        "左航与邓佳鑫是否存在2020年私人关系的正式声明？",
        "匿名爆料原始发帖人与原始证据是否可追溯？",
        "第三方权威媒体是否曾对该传闻进行过署名调查报道？",
    ]

    evidence_bucket_extra = {
        "evidence": {
            cid: [e for e in evidences_flat if e.claim_id == cid]
            for cid in ("c1", "c2", "c3", "c4")
        },
    }

    return FactCheckReport(
        original_text=("TF家族三代成员左航和邓佳鑫曾经是情侣，2020年在恋爱中。" + "B" * 20),
        overall_verdict="证据不足",
        overall_summary=(
            "自然人身份类主张（左航、邓佳鑫为 TF家族三代公开训练生）已由 A/B 级多个独立来源"
            "一致证实；私生活传闻（曾为情侣 / 2020 年恋爱中）在排除同源转载与粉丝二次创作后，"
            "缺少 ≥2 个相互独立的高质量来源，且存在当事人所属公司的 A 级澄清声明，"
            "因此整体结论为「证据不足」。建议仅参考当事人声明与权威署名报道，避免转发匿名爆料。"
        ),
        claim_results=claim_results,
        timeline=timeline,
        propagation_risk=propagation_risk_str,
        unresolved_questions=unresolved,
        execution_log=[
            {"step": "init", "status": "success", "action": "初始化案例",
             "details": {"input": "用户输入文本"}},
            {"step": "decompose", "status": "success", "action": "拆解4条主张",
             "details": {"context": 2, "core": 1, "causal_or_detail": 1}},
            {"step": "search", "status": "success", "action": "第一轮检索完成",
             "details": {"round": 1}},
            {"step": "sufficiency", "status": "success",
             "action": "因证据不足，已启动第2轮补充检索（身份类查询补充官方关键词）",
             "details": {"reason": "私生活传闻缺少≥2个独立高质量来源",
                         "round": 2,
                         "queries_by_claim": {"c3": ["左航 邓佳鑫 声明", "左航 邓佳鑫 报道"],
                                              "c4": ["左航 邓佳鑫 2020 声明"]}}},
            {"step": "search", "status": "success", "action": "第二轮检索完成",
             "details": {"round": 2}},
            {"step": "evaluate", "status": "success", "action": "完成交叉验证",
             "details": {"verdict_counts": {"已证实": 2, "证据不足": 2}}},
            {"step": "report", "status": "success", "action": "生成最终报告",
             "details": {"overall": "证据不足"}},
        ],
        current_step="completed",
        completed_steps=["init", "decompose", "plan", "search",
                         "sufficiency", "evaluate", "report"],
        skipped_steps=["memory"],
        progress_percent=100,
        workflow_completed=True,
        workflow_error=None,
        generated_at=now,
        **plans_extra,           # 非模型显式字段：plan（list[VerificationPlan]）
        **evidence_bucket_extra, # 非模型显式字段：evidence（dict[str,list[Evidence]]，兼容 ctx.evidence_index）
    )


# ========================= helpers =========================
def _iter_paragraphs_and_cells(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _text_of(paragraphs) -> str:
    out: list[str] = []
    for p in paragraphs:
        if isinstance(p, Paragraph):
            out.append(p.text)
        elif isinstance(p, _Cell):
            out.append(p.text)
    return "\n".join(out)


def _all_text(doc: Document) -> str:
    parts: list[str] = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            parts.append(p.text)
        for p in section.footer.paragraphs:
            parts.append(p.text)
    for p in _iter_paragraphs_and_cells(doc):
        parts.append(p.text)
    return "\n".join(parts)


def _has_real_hyperlink(doc: Document, url: str) -> bool:
    """检查文档中是否存在真正的 OOXML 超链接关系（不是纯文本假装的链接）。"""
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype and rel.target_ref == url:
            # 同时确认它被正文引用
            for p in _iter_paragraphs_and_cells(doc):
                for r in p.runs:
                    pass  # 不在 runs 层级查
            # 查找 body 中 <w:hyperlink>
            xpath = ".//w:hyperlink"
            for p in _iter_paragraphs_and_cells(doc):
                # paragraph._p
                elem = getattr(p, "_p", None)
                if elem is None:
                    continue
                for h in elem.findall(qn("w:hyperlink")):
                    rid = h.get(qn("r:id"))
                    if rid:
                        try:
                            target = doc.part.rels[rid].target_ref
                        except Exception:  # noqa: BLE001
                            continue
                        if target == url:
                            return True
    return False


def _no_fixed_row_height(doc: Document) -> bool:
    """所有表格行没有设置 trHeight 的固定 hRule="exact"。"""
    for tbl in doc.tables:
        for row in tbl.rows:
            tr = row._tr
            trPr = tr.find(qn("w:trPr"))
            if trPr is not None:
                trHeight = trPr.find(qn("w:trHeight"))
                if trHeight is not None:
                    rule = trHeight.get(qn("w:hRule"))
                    if rule == "exact":
                        return False
    return True


# ========================= tests =========================
def test_docx_opens_and_has_sections_with_timeline() -> None:
    report = _make_report(include_timeline=True)
    data = build_fact_check_docx(report)
    assert isinstance(data, (bytes, bytearray)) and len(data) > 0
    doc = Document(__import__("io").BytesIO(data))

    txt = _all_text(doc)
    # 章节存在
    for heading in [
        "溯真｜新闻事实核查报告",
        "核查摘要",
        "主张核查总表",
        "逐项核查详情",
        "证据与来源",
        "事件时间线",
        "传播风险",
        "仍待核实的问题",
        "核查方法与免责声明",
        "完整来源目录",
    ]:
        assert heading in txt, f"缺少章节标题：{heading}"

    # 核心结论语义：背景身份证实，私生活传闻证据不足
    assert "左航是TF家族三代成员" in txt
    assert "邓佳鑫是TF家族三代成员" in txt
    # 高风险传闻不能因身份证实而属实（必须有"证据不足"出现在对应段落；检查存在"已证实"也存在"证据不足"）
    assert "已证实" in txt
    assert "证据不足" in txt

    # 时间线
    assert "2020-12-25" in txt or "匿名粉丝发帖提及传闻" in txt

    # 不包含密钥 / 环境变量名的值（我们把伪造密钥塞到环境变量里的场景这里直接检查纯文本）
    for bad in [
        "LLM_API_KEY", "TAVILY_API_KEY", "sk-", "tk-",
    ]:
        assert bad not in txt, f"文档中出现不允许出现的字符串：{bad}"

    # 没有固定行高导致文字被截断
    assert _no_fixed_row_height(doc)

    # 存在真实超链接（证据链接能点）
    for url in [
        "https://tf.example.com/zuohang",
        "https://news.example.com/tf-3rd",
        "https://news.example.com/deny",
    ]:
        assert _has_real_hyperlink(doc, url), f"文档中缺少真实可点击超链接：{url}"

    # 页眉/页脚
    headers_joined = "\n".join(p.text for s in doc.sections for p in s.header.paragraphs)
    footers_joined = "\n".join(p.text for s in doc.sections for p in s.footer.paragraphs)
    assert "溯真｜事实核查报告" in headers_joined
    assert "生成时间" in footers_joined or "页码" in footers_joined


def test_docx_no_timeline_message() -> None:
    report = _make_report(include_timeline=False)
    data = build_fact_check_docx(report)
    doc = Document(__import__("io").BytesIO(data))
    txt = _all_text(doc)
    assert "现有证据不足以构建可靠时间线" in txt


def test_docx_no_secrets_in_bytes_or_text(monkeypatch) -> None:
    """即使当前进程环境变量包含密钥，导出的文档也绝不包含这些敏感值。"""
    monkeypatch.setenv("LLM_API_KEY", "sk-MOCKSECRETKEYFORTEST123456")
    monkeypatch.setenv("TAVILY_API_KEY", "tvly-MOCKSECRETKEYFORTEST987654")
    monkeypatch.setenv("OPENAI_API_KEY", "sk-MOCKOPENAI9999")

    report = _make_report(include_timeline=True)
    data = build_fact_check_docx(report)
    # 原始字节里也不能出现（避免 base64/utf8 泄漏）
    raw = bytes(data)
    for token in [b"sk-MOCKSECRETKEYFORTEST123456", b"tvly-MOCKSECRETKEYFORTEST987654", b"sk-MOCKOPENAI9999"]:
        assert token not in raw
    doc = Document(__import__("io").BytesIO(raw))
    txt = _all_text(doc)
    for bad in ["MOCKSECRETKEYFORTEST", "MOCKOPENAI9999", "sk-MOCK", "tvly-MOCK", "LLM_API_KEY", "TAVILY_API_KEY"]:
        assert bad not in txt


def test_docx_risk_colors_exist() -> None:
    """检查表格中的结论颜色映射确实写在 run 属性里（通过段落文本 + XML 属性推断）。"""
    report = _make_report(include_timeline=True)
    data = build_fact_check_docx(report)
    doc = Document(__import__("io").BytesIO(data))
    # 在主张总表里找到 "已证实"/"证据不足"；检查至少存在对应的颜色 XML（不是断言具体哪个段落，只断言 XML 里有这些颜色值）
    hexes = []
    for p in _iter_paragraphs_and_cells(doc):
        for run in p.runs:
            rPr = run._element.find(qn("w:rPr"))
            if rPr is None:
                continue
            color = rPr.find(qn("w:color"))
            if color is not None:
                hexes.append(color.get(qn("w:val")))
    # 绿色/黄色/深蓝色至少出现一次（已证实/证据不足/主色）——集合包含 COLOR_SUCCESS / COLOR_WARN / COLOR_MAIN 的精确 hex
    assert any(v in {"2E8B57", "1F7A3C", "2E7D32", "1E7B32"} for v in hexes)  # 绿（含精确 COLOR_SUCCESS=1E7B32）
    assert any(v in {"D4A017", "B8860B", "D2691E", "F4C430"} for v in hexes)  # 黄/琥珀（含精确 COLOR_WARN=B8860B）
    assert any(v in {"1F4E79", "1F4E78", "003366", "1E3A8A", "1F4E79", "0B3D91"} for v in hexes)  # 深蓝（含精确 COLOR_MAIN=0B3D91）
