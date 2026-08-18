"""Word (DOCX) 事实核查简报导出器。

只使用结构化 FactCheckReport 字段进行渲染。
- 不包含 API 密钥、系统提示词、隐藏推理或内部异常日志
- 不保存整页网页正文
- 输出为 BytesIO 中的 bytes，使用 A4 纵向 + 中文等线/宋体
"""

from __future__ import annotations

import datetime as _dt
import io
from dataclasses import dataclass
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from src.models import FactCheckReport, ClaimResult, Evidence

# ===== 版式参数（A4 中文正式报告）=====
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_CM = 2.2
MARGIN_BOTTOM_CM = 2.2
MARGIN_LEFT_CM = 2.4
MARGIN_RIGHT_CM = 2.4

FONT_TITLE_CN = "等线"  # 标题优先等线，也可改成"微软雅黑"
FONT_HEADING_CN = "等线"
FONT_BODY_CN = "宋体"

FONT_SIZE_TITLE = Pt(20)
FONT_SIZE_SUBTITLE = Pt(12)
FONT_SIZE_HEADING1 = Pt(16)
FONT_SIZE_HEADING2 = Pt(13)
FONT_SIZE_BODY = Pt(10.5)
LINE_SPACING = 1.3

# 主色：深蓝；语义色：绿=已证实 黄=证据不足 红=已证伪/高风险 灰=背景说明
COLOR_MAIN = RGBColor(0x0B, 0x3D, 0x91)      # 深蓝
COLOR_SUCCESS = RGBColor(0x1E, 0x7B, 0x32)   # 绿
COLOR_WARN = RGBColor(0xB8, 0x86, 0x0B)      # 黄（较深，保持可读）
COLOR_DANGER = RGBColor(0xA5, 0x2A, 0x2A)    # 红
COLOR_GREY = RGBColor(0x59, 0x59, 0x59)      # 灰
COLOR_LIGHT_BG = "D9E2F3"                    # 表头浅蓝填充
COLOR_HEADER_DARK = "0B3D91"                 # 页眉深蓝文字


CLAIM_ROLE_LABEL = {
    "context": "背景/身份前提",
    "core": "核心事件主张",
    "causal_or_detail": "原因/时间/数量细节",
}


# ===== 中文字体辅助 =====

def _set_run_font(run, size_pt, bold=False, color=None, east_asia=FONT_BODY_CN, latin="Calibri"):
    """同时设置 ascii / hAnsi / eastAsia 字体，避免 Word 里中文退化。"""
    run.font.name = latin
    run.font.size = size_pt
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:eastAsia"), east_asia)
    return run


def _add_paragraph(
    doc: Document,
    text: str = "",
    size=FONT_SIZE_BODY,
    bold=False,
    color=None,
    east_asia=FONT_BODY_CN,
    align=None,
    style=None,
    line_spacing=None,
) -> Paragraph:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if text:
        r = p.add_run(text)
        _set_run_font(r, size, bold, color, east_asia=east_asia)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = LINE_SPACING if line_spacing is None else line_spacing
    fmt.space_after = Pt(4)
    fmt.space_before = Pt(0)
    return p


def _apply_heading_style(p: Paragraph, level: int, color=COLOR_MAIN):
    """对 Word 内置 Heading 样式的中文 heading 补充中文字体设置。"""
    size = {1: FONT_SIZE_HEADING1, 2: FONT_SIZE_HEADING2}.get(level, FONT_SIZE_HEADING2)
    for run in p.runs:
        _set_run_font(run, size, bold=True, color=color, east_asia=FONT_HEADING_CN)
    return p


def _set_cell_bg(cell: _Cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _style_cell(cell: _Cell, bold=False, color=None, align=None, size=FONT_SIZE_BODY, no_fixed_row=True):
    """对单元格所有 run 设置中文正文字体，并对表格段落应用 1.3 倍行距。"""
    for para in cell.paragraphs:
        para.paragraph_format.line_spacing = LINE_SPACING
        para.paragraph_format.space_after = Pt(2)
        if align is not None:
            para.alignment = align
        for run in para.runs:
            _set_run_font(run, size, bold=bold, color=color, east_asia=FONT_BODY_CN)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # 禁止固定行高（防止文字截断）：默认 Word 表格允许换行；此处只保证段落有行高。
    _ = no_fixed_row
    return cell


def _write_cell(cell: _Cell, text, bold=False, color=None, align=None, size=FONT_SIZE_BODY):
    # 清空原有文字再写
    for p in list(cell.paragraphs):
        for r in list(p.runs):
            r.text = ""
        if cell.paragraphs[0] is not p:
            p._element.getparent().remove(p._element)
    if cell.paragraphs:
        cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    if text is None:
        text = ""
    if text == "":
        _style_cell(cell, bold=bold, color=color, align=align, size=size)
        return
    r = p.add_run(str(text))
    _set_run_font(r, size, bold=bold, color=color, east_asia=FONT_BODY_CN)
    _style_cell(cell, bold=bold, color=color, align=align, size=size)


def _add_hyperlink(paragraph: Paragraph, text: str, url: str, color=COLOR_MAIN, bold=False):
    """添加真实 Word 超链接（可点击 + 存在于 hyperlinks_relationships 中）。"""
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:rId"), r_id)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rPr.append(color_el)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    _ = bold  # 这里用 Hyperlink 样式，不额外粗体
    return r_id


def _verdict_color(verdict: str) -> RGBColor:
    v = (verdict or "").lower()
    if "已证伪" in verdict or "已被驳回" in verdict or "被证实为虚假" in verdict or "refute" in v or v == "false":
        return COLOR_DANGER
    if "已证实" in verdict or "已确认" in verdict or "基本属实" in verdict or "证实" in verdict:
        return COLOR_SUCCESS
    return COLOR_WARN


def _risk_color(level: str) -> RGBColor:
    return {"high": COLOR_DANGER, "medium": COLOR_WARN}.get(level or "", COLOR_GREY)


@dataclass
class _RenderContext:
    doc: "Document"
    report: FactCheckReport
    generated_at: str
    report_number: str

    # ===== 与当前 FactCheckReport schema 对齐的兼容视图（避免 AttributeError）=====
    @property
    def evidence_index(self) -> list:
        """DOCX 需要 Evidence 的扁平列表；FactCheckReport.evidence 是 dict[str,list[Evidence]]。"""
        evs: list = []
        bucket = getattr(self.report, "evidence", None) or {}
        if isinstance(bucket, dict):
            for vs in bucket.values():
                if vs:
                    evs.extend(vs)
        elif isinstance(bucket, list):
            evs.extend(bucket)
        return evs

    @property
    def evidence_pool(self) -> list:
        return self.evidence_index

    @property
    def overall_confidence(self) -> str:
        """兼容 ctx.report.overall_confidence（旧字段）。"""
        # 如果没有单条结果，返回未评估；否则简单取平均置信度
        results = self.report.claim_results or []
        if not results:
            return "未评估"
        vals = [float(getattr(r, "confidence", 0) or 0) for r in results]
        avg = sum(vals) / len(vals)
        return f"{avg:.0%}" if avg <= 1 else f"{avg:.1f}"

    @property
    def metadata(self) -> dict:
        """兼容 ctx.report.metadata（旧字段），返回 mode + overall_risk。"""
        # FactCheckReport 现在没有 metadata 字段；用 execution_log + propagation_risk 推断
        propagation_risk = (getattr(self.report, "propagation_risk", "") or "").strip()
        # 从 propagation_risk 中提取等级字符串，如 "中等" → "medium" / "高" → "high"
        level = "low"
        for tag, key in (("高", "high"), ("严重", "high"), ("中", "medium"), ("低", "low")):
            if tag in propagation_risk:
                level = key
                break
        # 用 workflow 的 mode 映射：FactCheckReport 没有 mode，但通常执行日志里包含 mode_label
        mode = "未知"
        # 尝试从 completed_steps 等推断：这里用 report 上的兜底
        return {"overall_risk": level, "mode": mode}

    @property
    def summary(self) -> str:
        return getattr(self.report, "overall_summary", "") or ""

    @property
    def claims(self) -> list:
        return list(getattr(self.report, "claims", []) or [])


# ===== 章节 =====

def _render_masthead(ctx: _RenderContext):
    """memo_masthead：标题 + 副标题 + 元数据表格。"""
    doc = ctx.doc
    p = _add_paragraph(doc, "溯真｜新闻事实核查报告", size=FONT_SIZE_TITLE, bold=True,
                       color=COLOR_MAIN, east_asia=FONT_TITLE_CN, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(2)

    # 副标题：原始待核查文本的简短摘要（首 60 字）
    subtitle = (ctx.report.original_text or "").strip().replace("\n", " ")
    if len(subtitle) > 60:
        subtitle = subtitle[:60] + "…"
    _add_paragraph(doc, f"待核查摘要：{subtitle or '（无）'}", size=FONT_SIZE_SUBTITLE,
                   color=COLOR_GREY, east_asia=FONT_BODY_CN)

    # 元数据表（两列：键值）
    total_ev = len(ctx.report.evidence_index or [])
    unique_sources = len({e.source_url for e in (ctx.report.evidence_index or []) if getattr(e, "source_url", None)})
    claim_count = len(ctx.report.claim_results or [])

    meta_rows = [
        ("报告编号", ctx.report_number),
        ("生成时间", ctx.generated_at),
        ("核查模式", (ctx.report.metadata or {}).get("mode", "未知")),
        ("总体结论", ctx.report.overall_verdict or "未生成"),
        ("主张数量", f"{claim_count} 条"),
        ("有效证据数量", f"{total_ev} 条"),
        ("独立来源数量", f"{unique_sources} 个"),
    ]

    t = doc.add_table(rows=len(meta_rows), cols=2)
    try:
        t.columns[0].width = Cm(3.5)
        t.columns[1].width = Cm(PAGE_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM - 3.5)
    except Exception:
        pass

    for i, (k, v) in enumerate(meta_rows):
        key_cell = t.cell(i, 0)
        val_cell = t.cell(i, 1)
        _write_cell(key_cell, k, bold=True, color=COLOR_MAIN, size=FONT_SIZE_BODY)
        _set_cell_bg(key_cell, COLOR_LIGHT_BG)
        _write_cell(val_cell, v, size=FONT_SIZE_BODY)
        # 总体结论单独着色
        if k == "总体结论":
            for p in val_cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = _verdict_color(ctx.report.overall_verdict)
    _add_paragraph(doc, "")


def _add_heading(ctx, text, level=1):
    style = {1: "Heading 1", 2: "Heading 2"}.get(level, "Heading 3")
    p = ctx.doc.add_paragraph(style=style)
    r = p.add_run(text)
    color = COLOR_MAIN
    _set_run_font(r,
                  {1: FONT_SIZE_HEADING1, 2: FONT_SIZE_HEADING2}.get(level, FONT_SIZE_BODY),
                  bold=True, color=color, east_asia=FONT_HEADING_CN)
    _apply_heading_style(p, level, color=color)
    return p


def _render_exec_summary(ctx: _RenderContext):
    """1. 核查摘要。"""
    _add_heading(ctx, "1. 核查摘要", level=1)

    doc = ctx.doc
    verdict = ctx.report.overall_verdict or "未生成"
    confidence = ctx.report.overall_confidence or "未评估"
    risk = (ctx.report.metadata or {}).get("overall_risk", "low")

    # 彩色结论框：用一行一格表格实现
    box_t = doc.add_table(rows=1, cols=1)
    cell = box_t.cell(0, 0)
    _set_cell_bg(cell, "F2F5FC")
    box_p = cell.paragraphs[0]
    r1 = box_p.add_run("总体结论：")
    _set_run_font(r1, FONT_SIZE_BODY, bold=True, color=COLOR_MAIN)
    r2 = box_p.add_run(verdict)
    _set_run_font(r2, FONT_SIZE_BODY, bold=True, color=_verdict_color(verdict))
    box_p2 = cell.add_paragraph()
    r3 = box_p2.add_run(f"总体置信度：{confidence}   |   总体风险等级：{risk.upper()}")
    _set_run_font(r3, FONT_SIZE_BODY, color=COLOR_GREY)

    # 2-4 句话核心判断
    summary_lines = _build_summary_sentences(ctx)
    for line in summary_lines:
        _add_paragraph(doc, "• " + line, size=FONT_SIZE_BODY, color=COLOR_GREY)


def _build_summary_sentences(ctx: _RenderContext) -> list[str]:
    crs: list[ClaimResult] = ctx.report.claim_results or []
    if not crs:
        return ["无可用主张，无法生成核心判断。"]
    lines: list[str] = []
    contexts = [c for c in crs if (c.claim or {}).get("claim_role") == "context"]
    cores = [c for c in crs if (c.claim or {}).get("claim_role") == "core"]
    details = [c for c in crs if (c.claim or {}).get("claim_role") == "causal_or_detail"]

    if contexts:
        verified = sum(1 for c in contexts if "已证实" in (c.verdict or ""))
        lines.append(f"共拆解 {len(contexts)} 条背景/身份前提主张，其中 {verified} 条已被可靠来源证实。")
    if cores:
        for c in cores[:1]:
            verdict = c.verdict or "未判定"
            lines.append(f"核心事件主张「{(c.claim or {}).get('text','')[:40]}…」判定为：{verdict}。")
    if details:
        lines.append(f"时间/数量/原因类附加细节主张共 {len(details)} 条，已逐条独立核查。")

    if any("证据不足" in (c.verdict or "") for c in crs) and not any("高风险" in (c.claim or {}).get("sensitive_reason","") for c in crs):
        lines.append("仍存在关键证据缺口，建议关注文末「仍待核实的问题」章节。")
    # 明确提醒：背景身份已证实不代表私生活传闻属实
    if contexts and cores:
        lines.append("背景/身份类主张与私生活传闻主张分别独立判断；前者证实不自动推高后者可信度。")
    # 补齐 2-4 条
    while len(lines) < 2:
        lines.append("具体证据、来源与判定理由见后续章节。")
    return lines[:4]


def _render_claim_summary_table(ctx: _RenderContext):
    """2. 主张核查总表。"""
    _add_heading(ctx, "2. 主张核查总表", level=1)
    crs = ctx.report.claim_results or []
    if not crs:
        _add_paragraph(ctx.doc, "（无主张）", color=COLOR_GREY)
        return

    headers = ["编号", "主张", "主张角色", "结论", "置信度", "风险等级"]
    # 列宽：1/8/2.5/2.5/1.8/1.8 （总约 17.6cm，适配 A4 左右页边距）
    widths_cm = [1.2, 7.8, 2.6, 2.5, 1.7, 1.7]

    t: Table = ctx.doc.add_table(rows=1 + len(crs), cols=len(headers))
    for i, w in enumerate(widths_cm):
        try:
            t.columns[i].width = Cm(w)
        except Exception:
            pass

    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        _write_cell(cell, h, bold=True, color=COLOR_MAIN,
                    align=WD_ALIGN_PARAGRAPH.CENTER, size=FONT_SIZE_BODY)
        _set_cell_bg(cell, COLOR_LIGHT_BG)
        _style_cell(cell, bold=True, color=COLOR_MAIN)

    for i, cr in enumerate(crs, start=1):
        claim = cr.claim or {}
        role = CLAIM_ROLE_LABEL.get(claim.get("claim_role", ""), claim.get("claim_role", "未知"))
        verdict = cr.verdict or "未判定"
        conf = cr.confidence or ""
        risk = claim.get("risk_level", "low")
        risk_txt = {"high": "高", "medium": "中", "low": "低"}.get(risk or "", risk or "")

        cells = [
            f"{i}",
            (claim.get("text") or "")[:80],
            role,
            verdict,
            conf,
            risk_txt,
        ]
        for j, v in enumerate(cells):
            cell = t.cell(i, j)
            align = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2, 3, 4, 5) else WD_ALIGN_PARAGRAPH.LEFT
            _write_cell(cell, v, size=FONT_SIZE_BODY, align=align)
            # 结论 / 风险列着色
            if j == 3:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = _verdict_color(verdict)
            if j == 5:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = _risk_color(risk)
            _style_cell(cell, align=align)


def _render_claim_details(ctx: _RenderContext):
    """3. 逐项核查详情。"""
    _add_heading(ctx, "3. 逐项核查详情", level=1)
    crs = ctx.report.claim_results or []
    ev_map = {e.evidence_id: e for e in (ctx.report.evidence_index or [])}

    for idx, cr in enumerate(crs, start=1):
        claim = cr.claim or {}
        _add_heading(ctx, f"3.{idx}  主张 C{idx}：{(claim.get('text') or '')[:60]}", level=2)

        role = CLAIM_ROLE_LABEL.get(claim.get("claim_role", ""), claim.get("claim_role", "未知"))
        risk_txt = {"high": "高", "medium": "中", "low": "低"}.get(claim.get("risk_level", "low"), "低")

        rows = [
            ("主张原文", claim.get("text") or ""),
            ("主张角色", role),
            ("结论", cr.verdict or "未判定"),
            ("置信度", cr.confidence or ""),
            ("风险等级", risk_txt),
            ("判断理由", cr.justification or ""),
            ("缺失信息", cr.missing_info or ""),
            ("所用证据编号", ", ".join(cr.used_evidence_ids or []) or "（无）"),
        ]
        t = ctx.doc.add_table(rows=len(rows), cols=2)
        try:
            t.columns[0].width = Cm(2.5)
            t.columns[1].width = Cm(PAGE_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM - 2.5)
        except Exception:
            pass
        for i, (k, v) in enumerate(rows):
            key_cell = t.cell(i, 0)
            val_cell = t.cell(i, 1)
            _write_cell(key_cell, k, bold=True, color=COLOR_MAIN, size=FONT_SIZE_BODY)
            _set_cell_bg(key_cell, COLOR_LIGHT_BG)
            _write_cell(val_cell, v, size=FONT_SIZE_BODY)
            if k == "结论":
                for p in val_cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = _verdict_color(cr.verdict or "")
            _style_cell(key_cell)
            _style_cell(val_cell)


def _truncate_for_excerpt(text: str, maxlen=300) -> str:
    if not text:
        return ""
    t = (text or "").strip()
    if len(t) <= maxlen:
        return t
    return t[:maxlen].rstrip() + "…"


def _render_evidence(ctx: _RenderContext):
    """4. 证据与来源（含可点击真实链接）。"""
    _add_heading(ctx, "4. 证据与来源", level=1)
    evs: list[Evidence] = ctx.report.evidence_index or []
    if not evs:
        _add_paragraph(ctx.doc, "（无证据）", color=COLOR_GREY)
        return

    stance_label = {"supports": "支持", "refutes": "反驳", "context": "背景/上下文", "irrelevant": "无关"}
    for idx, ev in enumerate(evs, start=1):
        _add_heading(ctx, f"证据 [{idx}]  {ev.evidence_id}  {_truncate_for_excerpt(ev.source_title or ev.source_url or '', 60)}", level=2)

        rows = [
            ("来源标题", ev.source_title or ""),
            ("发布机构", ev.publisher or ev.source_domain or ""),
            ("来源等级", f"{ev.source_grade.upper()} 级" if getattr(ev, "source_grade", None) else ""),
            ("支持或反驳", stance_label.get(getattr(ev, "stance", ""), getattr(ev, "stance", "") or "")),
            ("发布日期", getattr(ev, "published_at", "") or ""),
        ]
        t = ctx.doc.add_table(rows=len(rows), cols=2)
        try:
            t.columns[0].width = Cm(2.5)
            t.columns[1].width = Cm(PAGE_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM - 2.5)
        except Exception:
            pass
        for i, (k, v) in enumerate(rows):
            key_cell = t.cell(i, 0)
            val_cell = t.cell(i, 1)
            _write_cell(key_cell, k, bold=True, color=COLOR_MAIN)
            _set_cell_bg(key_cell, COLOR_LIGHT_BG)
            # 标题/发布机构如果有 url 则叠加超链接
            if k == "来源标题" and ev.source_url:
                if val_cell.paragraphs:
                    for r in list(val_cell.paragraphs[0].runs):
                        r.text = ""
                    _add_hyperlink(val_cell.paragraphs[0], v or ev.source_url, ev.source_url, color=COLOR_MAIN)
            else:
                _write_cell(val_cell, v)
            _style_cell(key_cell)
            _style_cell(val_cell)

        # 有效证据摘录：≤300 字
        excerpt = _truncate_for_excerpt(ev.relevant_excerpt or "", 300)
        p_excerpt = _add_paragraph(ctx.doc, "")
        r1 = p_excerpt.add_run("有效证据摘录（≤300字）：")
        _set_run_font(r1, FONT_SIZE_BODY, bold=True, color=COLOR_MAIN)
        if excerpt:
            p = _add_paragraph(ctx.doc, f"“{excerpt}”", size=FONT_SIZE_BODY, color=COLOR_GREY)
            p.paragraph_format.left_indent = Cm(0.5)
        else:
            _add_paragraph(ctx.doc, "（未提供有效摘录）", color=COLOR_GREY)

        if ev.source_url:
            p_link = _add_paragraph(ctx.doc, "")
            r2 = p_link.add_run("来源链接：")
            _set_run_font(r2, FONT_SIZE_BODY, bold=True, color=COLOR_MAIN)
            _add_hyperlink(p_link, ev.source_url, ev.source_url, color=COLOR_MAIN)


def _render_timeline(ctx: _RenderContext):
    """5. 事件时间线。"""
    _add_heading(ctx, "5. 事件时间线", level=1)
    evs = [e for e in (ctx.report.evidence_index or []) if getattr(e, "published_at", None)]

    def _sort_key(e: Evidence):
        try:
            return _dt.datetime.fromisoformat(str(e.published_at))
        except Exception:
            return _dt.datetime.max
    evs_sorted = sorted(evs, key=_sort_key)

    if not evs_sorted:
        _add_paragraph(ctx.doc, "现有证据不足以构建可靠时间线（无可靠发布日期）。", color=COLOR_GREY)
        return

    t = ctx.doc.add_table(rows=1 + len(evs_sorted), cols=3)
    try:
        t.columns[0].width = Cm(3.0)
        t.columns[1].width = Cm(9.2)
        t.columns[2].width = Cm(4.4)
    except Exception:
        pass
    headers = ["发布日期", "事件 / 证据摘要", "来源"]
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        _write_cell(c, h, bold=True, color=COLOR_MAIN, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(c, COLOR_LIGHT_BG)
        _style_cell(c, bold=True)

    for i, ev in enumerate(evs_sorted, start=1):
        cells = [
            str(ev.published_at),
            _truncate_for_excerpt(ev.relevant_excerpt or ev.source_title or "", 120),
            _truncate_for_excerpt(ev.publisher or ev.source_domain or "", 40),
        ]
        for j, v in enumerate(cells):
            c = t.cell(i, j)
            _write_cell(c, v, align=WD_ALIGN_PARAGRAPH.LEFT if j in (1, 2) else WD_ALIGN_PARAGRAPH.CENTER)
            _style_cell(c)


def _render_risk(ctx: _RenderContext):
    """6. 传播风险。"""
    _add_heading(ctx, "6. 传播风险", level=1)
    crs = ctx.report.claim_results or []
    high_claims = [c for c in crs if (c.claim or {}).get("risk_level") == "high"]
    lines: list[str] = []
    if high_claims:
        lines.append(f"共识别 {len(high_claims)} 条高风险主张（涉及隐私、名誉、公共安全等）。")
        for c in high_claims:
            txt = (c.claim or {}).get("text", "")[:70]
            lines.append(f"- 「{txt}」：判定为 {(c.verdict or '未判定')}，需谨慎转发。")
    else:
        lines.append("未识别高风险主张。")

    overall = (ctx.report.metadata or {}).get("overall_risk", "low")
    lines.append(f"本报告总体传播风险评级：{overall.upper()}")
    # 私生活高风险特别提醒（独立变量，避免在复杂函数参数中出现字符串拼接/引号错误）
    privacy_notice = (
        "特别提示：涉及自然人私生活的传闻，在缺乏当事人声明"
        "或多个相互独立的高质量来源时，结论通常不高于「证据不足」，"
        "不得以粉丝讨论、匿名爆料或剪辑视频作为证实依据。"
    )
    if any("情侣" in (c.claim or {}).get("text", "") or "恋爱" in (c.claim or {}).get("text", "") for c in crs):
        lines.append(privacy_notice)
    for line in lines:
        _add_paragraph(ctx.doc, line, size=FONT_SIZE_BODY, color=COLOR_DANGER if "高" in line else COLOR_GREY)


def _render_open_questions(ctx: _RenderContext):
    """7. 仍待核实的问题。"""
    _add_heading(ctx, "7. 仍待核实的问题", level=1)
    crs = ctx.report.claim_results or []
    items: list[str] = []
    for c in crs:
        if c.missing_info and (c.missing_info or "").strip():
            items.append(f"- 主张「{(c.claim or {}).get('text','')[:60]}」缺失：{c.missing_info.strip()}")
    if not items:
        items = ["（当前报告未列出仍待核实的问题。）"]
    for line in items:
        _add_paragraph(ctx.doc, line, size=FONT_SIZE_BODY, color=COLOR_GREY)


def _render_methodology(ctx: _RenderContext):
    """8. 核查方法与免责声明。"""
    _add_heading(ctx, "8. 核查方法与免责声明", level=1)
    lines = [
        "• 本报告使用结构化事实核查流程：主张拆解 → 核查计划 → 自动检索 → 证据充分性评估（不足时自动补充检索） → 独立判定 → 结果整理。",
        "• 主张独立判定原则：背景身份类、核心事件类、时间/数量细节类主张分别独立判断；背景事实被证实不自动提高其他传闻的可信度。",
        "• 私生活/名誉类主张执行更严格的证据门槛：需当事人公开声明、可靠采访或多个独立高质量来源；粉丝讨论、匿名爆料、剪辑视频仅视为 D 级线索；多篇转载同一爆料不视为多个独立来源。",
        "• 来源等级标准：A=官方一手；B=权威署名媒体；C=一般媒体/规范转载；D=自媒体/社交平台；E=匿名或无法核实。",
        "• 免责声明：本报告基于公开可检索网页在核查时点的快照生成；如遇网页撤稿、更正或后续新证据，需重新启动核查。本报告结论不构成法律意见。",
    ]
    for line in lines:
        _add_paragraph(ctx.doc, line, size=FONT_SIZE_BODY, color=COLOR_GREY)


def _render_source_directory(ctx: _RenderContext):
    """9. 完整来源目录：[1]编号，标题+网址可点击，与正文一致。"""
    _add_heading(ctx, "9. 完整来源目录", level=1)
    evs = ctx.report.evidence_index or []
    if not evs:
        _add_paragraph(ctx.doc, "（无来源）", color=COLOR_GREY)
        return
    for i, ev in enumerate(evs, start=1):
        p = _add_paragraph(ctx.doc, "", size=FONT_SIZE_BODY)
        r1 = p.add_run(f"[{i}] ")
        _set_run_font(r1, FONT_SIZE_BODY, bold=True, color=COLOR_MAIN)
        title = (ev.source_title or "").strip() or ev.source_url or ""
        if ev.source_url:
            _add_hyperlink(p, title or ev.source_url, ev.source_url, color=COLOR_MAIN)
        else:
            r2 = p.add_run(title or "(无链接)")
            _set_run_font(r2, FONT_SIZE_BODY, color=COLOR_GREY)
        if ev.source_url:
            p2 = _add_paragraph(ctx.doc, "", size=FONT_SIZE_BODY)
            r3 = p2.add_run("    URL：")
            _set_run_font(r3, FONT_SIZE_BODY, color=COLOR_GREY)
            _add_hyperlink(p2, ev.source_url, ev.source_url, color=COLOR_MAIN)


# ===== 页眉页脚 =====

def _setup_headers_footers(section, generated_at: str):
    """页眉：溯真｜事实核查报告；页脚：页码 + 生成时间。"""
    section.different_first_page_header_footer = False

    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run("溯真｜事实核查报告")
    _set_run_font(r, Pt(9), bold=False, color=COLOR_MAIN, east_asia=FONT_TITLE_CN)

    # 页脚：左侧生成时间；右侧 PAGE / NUMPAGES 域
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run(f"生成时间：{generated_at}    |    第 ")
    _set_run_font(r1, Pt(9), color=COLOR_GREY)
    from docx.oxml import OxmlElement
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    r_run = OxmlElement("w:r")
    r_run.append(fldChar1)
    r_run.append(instrText)
    r_run.append(fldChar2)
    r_run.append(fldChar3)
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{COLOR_GREY[0]:02X}{COLOR_GREY[1]:02X}{COLOR_GREY[2]:02X}")
    rPr.append(color_el)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")  # 9pt = 18 half-pt
    rPr.append(sz)
    r_run.insert(0, rPr)
    fp._p.append(r_run)

    r2 = fp.add_run(" 页  |  共 ")
    _set_run_font(r2, Pt(9), color=COLOR_GREY)

    fldChar1b = OxmlElement("w:fldChar")
    fldChar1b.set(qn("w:fldCharType"), "begin")
    instrTextb = OxmlElement("w:instrText")
    instrTextb.set(qn("xml:space"), "preserve")
    instrTextb.text = "NUMPAGES"
    fldChar2b = OxmlElement("w:fldChar")
    fldChar2b.set(qn("w:fldCharType"), "separate")
    fldChar3b = OxmlElement("w:fldChar")
    fldChar3b.set(qn("w:fldCharType"), "end")

    r_runb = OxmlElement("w:r")
    r_runb.append(fldChar1b)
    r_runb.append(instrTextb)
    r_runb.append(fldChar2b)
    r_runb.append(fldChar3b)
    rPrb = OxmlElement("w:rPr")
    color_el2 = OxmlElement("w:color")
    color_el2.set(qn("w:val"), f"{COLOR_GREY[0]:02X}{COLOR_GREY[1]:02X}{COLOR_GREY[2]:02X}")
    rPrb.append(color_el2)
    sz2 = OxmlElement("w:sz")
    sz2.set(qn("w:val"), "18")
    rPrb.append(sz2)
    r_runb.insert(0, rPrb)
    fp._p.append(r_runb)

    r3 = fp.add_run(" 页")
    _set_run_font(r3, Pt(9), color=COLOR_GREY)


# ===== 公共入口 =====

def build_fact_check_docx(report: FactCheckReport) -> bytes:
    """构建 Word 核查简报返回内存中 bytes（不保存磁盘）。"""
    doc = Document()

    # 基础样式（正文中文）
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = FONT_SIZE_BODY
    rPr = normal._element.get_or_add_rPr()
    from docx.oxml import OxmlElement as _OxmlElement
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = _OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_BODY_CN)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")

    # 页面尺寸：A4 纵向
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    # 生成时间 / 报告编号
    now = _dt.datetime.now()
    generated_at = now.strftime("%Y-%m-%d %H:%M:%S")
    report_number = f"SUZHEN-{now.strftime('%Y%m%d-%H%M%S')}"

    # 页眉页脚
    _setup_headers_footers(section, generated_at)

    ctx = _RenderContext(doc=doc, report=report, generated_at=generated_at, report_number=report_number)

    # 渲染全部章节
    _render_masthead(ctx)
    _render_exec_summary(ctx)
    _render_claim_summary_table(ctx)
    _render_claim_details(ctx)
    _render_evidence(ctx)
    _render_timeline(ctx)
    _render_risk(ctx)
    _render_open_questions(ctx)
    _render_methodology(ctx)
    _render_source_directory(ctx)

    # 写入 BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


__all__ = ["build_fact_check_docx"]
